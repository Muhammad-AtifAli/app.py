import streamlit as st
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from io import BytesIO
import re
import platform
import shutil


# =========================================================
# TESSERACT SETUP
# =========================================================
def setup_tesseract():
    """
    Sets Tesseract path only for Windows.
    On Streamlit Cloud/Linux, it uses system-installed tesseract.
    """
    system_name = platform.system()

    if system_name == "Windows":
        windows_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

        if shutil.which("tesseract"):
            pytesseract.pytesseract.tesseract_cmd = shutil.which("tesseract")
        else:
            pytesseract.pytesseract.tesseract_cmd = windows_path

    else:
        # Streamlit Cloud/Linux
        pytesseract.pytesseract.tesseract_cmd = "tesseract"


setup_tesseract()


# =========================================================
# STREAMLIT SETTINGS
# =========================================================
st.set_page_config(
    page_title="PDF Text Extractor",
    page_icon="📄",
    layout="wide"
)


# =========================================================
# DESIGN
# =========================================================
st.markdown(
    """
    <style>
    .main-title {
        text-align: center;
        color: #1F4E79;
        font-size: 42px;
        font-weight: bold;
        margin-bottom: 10px;
    }
    .sub-title {
        text-align: center;
        font-size: 18px;
        color: #555;
        margin-bottom: 30px;
    }
    .info-box {
        background-color: #F5F8FF;
        padding: 18px;
        border-radius: 12px;
        border: 1px solid #D9E6F2;
        margin-bottom: 20px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown(
    '<div class="main-title">📄 PDF Text Extractor App</div>',
    unsafe_allow_html=True
)

st.markdown(
    '<div class="sub-title">Extract text from searchable and scanned PDFs. Download as TXT or Word file.</div>',
    unsafe_allow_html=True
)


# =========================================================
# CLEANING FUNCTIONS
# =========================================================
def clean_text(text):
    """
    Cleans text for TXT, preview, and Word export.
    Removes illegal characters that break Word files.
    """
    if text is None:
        return ""

    text = text.replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0B\x0C\x0E-\x1F]", "", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)

    return text.strip()


# =========================================================
# CHECK TESSERACT
# =========================================================
def is_tesseract_available():
    """
    Checks whether Tesseract OCR is available.
    """
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def is_language_available(language_code):
    """
    Checks whether selected OCR language is installed.
    """
    try:
        available_langs = pytesseract.get_languages(config="")
        return language_code in available_langs
    except Exception:
        return False


# =========================================================
# SEARCHABLE PDF EXTRACTION
# =========================================================
def extract_searchable_pdf_text(pdf_bytes):
    extracted_text = ""

    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(pdf_document)

    progress_bar = st.progress(0)
    status_text = st.empty()

    for page_number in range(total_pages):
        page = pdf_document[page_number]
        page_text = page.get_text("text")

        extracted_text += f"\n\n================ Page {page_number + 1} ================\n\n"
        extracted_text += page_text

        progress = int(((page_number + 1) / total_pages) * 100)
        progress_bar.progress(progress)
        status_text.write(f"Processing page {page_number + 1} of {total_pages}")

    pdf_document.close()

    return clean_text(extracted_text)


# =========================================================
# OCR EXTRACTION
# =========================================================
def extract_scanned_pdf_text(pdf_bytes, dpi=200, language="eng"):
    extracted_text = ""

    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(pdf_document)

    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)

    progress_bar = st.progress(0)
    status_text = st.empty()

    for page_number in range(total_pages):
        page = pdf_document[page_number]

        pix = page.get_pixmap(matrix=matrix, alpha=False)

        image = Image.frombytes(
            "RGB",
            [pix.width, pix.height],
            pix.samples
        )

        page_text = pytesseract.image_to_string(image, lang=language)

        extracted_text += f"\n\n================ Page {page_number + 1} OCR ================\n\n"
        extracted_text += page_text

        progress = int(((page_number + 1) / total_pages) * 100)
        progress_bar.progress(progress)
        status_text.write(f"OCR processing page {page_number + 1} of {total_pages}")

    pdf_document.close()

    return clean_text(extracted_text)


# =========================================================
# AUTO MODE EXTRACTION
# =========================================================
def auto_extract_pdf_text(pdf_bytes, dpi=200, language="eng"):
    extracted_text = ""

    pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
    total_pages = len(pdf_document)

    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)

    progress_bar = st.progress(0)
    status_text = st.empty()

    searchable_pages = 0
    ocr_pages = 0

    for page_number in range(total_pages):
        page = pdf_document[page_number]

        page_text = page.get_text("text").strip()

        if len(page_text) > 30:
            searchable_pages += 1
            final_text = page_text
            mode = "Searchable Text"
        else:
            ocr_pages += 1

            pix = page.get_pixmap(matrix=matrix, alpha=False)

            image = Image.frombytes(
                "RGB",
                [pix.width, pix.height],
                pix.samples
            )

            final_text = pytesseract.image_to_string(image, lang=language)
            mode = "OCR"

        extracted_text += f"\n\n================ Page {page_number + 1} ({mode}) ================\n\n"
        extracted_text += final_text

        progress = int(((page_number + 1) / total_pages) * 100)
        progress_bar.progress(progress)
        status_text.write(f"Processing page {page_number + 1} of {total_pages}")

    pdf_document.close()

    message = f"Completed. Searchable pages: {searchable_pages}, OCR pages: {ocr_pages}"

    return clean_text(extracted_text), message


# =========================================================
# TXT FILE CREATION
# =========================================================
def create_txt_file(text):
    text = clean_text(text)

    txt_buffer = BytesIO()
    txt_buffer.write(text.encode("utf-8", errors="ignore"))
    txt_buffer.seek(0)

    return txt_buffer


# =========================================================
# WORD FILE CREATION
# =========================================================
def create_word_file(text):
    text = clean_text(text)

    document = Document()
    document.add_heading("Extracted PDF Text", level=1)

    page_parts = re.split(r"=+\s*Page\s+", text)

    page_counter = 0

    for part in page_parts:
        part = part.strip()

        if not part:
            continue

        page_counter += 1
        document.add_heading(f"Page {page_counter}", level=2)

        lines = part.split("\n")

        for line in lines:
            line = clean_text(line)

            if line:
                chunk_size = 2500

                for i in range(0, len(line), chunk_size):
                    document.add_paragraph(line[i:i + chunk_size])

    word_buffer = BytesIO()
    document.save(word_buffer)
    word_buffer.seek(0)

    return word_buffer


# =========================================================
# SIDEBAR
# =========================================================
st.sidebar.header("⚙️ Extraction Settings")

extraction_mode = st.sidebar.radio(
    "Select extraction mode",
    [
        "Auto Detect",
        "Searchable PDF Only",
        "Scanned PDF OCR Only"
    ]
)

output_type = st.sidebar.radio(
    "Select output file type",
    [
        "TXT File",
        "Word File"
    ]
)

ocr_dpi = st.sidebar.slider(
    "OCR Image Quality DPI",
    min_value=150,
    max_value=300,
    value=200,
    step=50
)

ocr_language = st.sidebar.selectbox(
    "OCR Language",
    [
        "eng",
        "urd"
    ],
    index=0
)

st.sidebar.info(
    "Use 'eng' for English PDFs. Use 'urd' only if Urdu OCR language is installed."
)


# =========================================================
# OCR STATUS
# =========================================================
tesseract_available = is_tesseract_available()

if tesseract_available:
    st.sidebar.success("Tesseract OCR is available.")
else:
    st.sidebar.error("Tesseract OCR is not available.")

selected_language_available = False

if tesseract_available:
    selected_language_available = is_language_available(ocr_language)

    if selected_language_available:
        st.sidebar.success(f"OCR language available: {ocr_language}")
    else:
        st.sidebar.warning(f"OCR language not available: {ocr_language}")


# =========================================================
# FILE UPLOADER
# =========================================================
uploaded_pdf = st.file_uploader(
    "Upload your PDF file",
    type=["pdf"]
)


# =========================================================
# MAIN LOGIC
# =========================================================
if uploaded_pdf is not None:
    st.success("PDF uploaded successfully.")

    pdf_bytes = uploaded_pdf.getvalue()

    try:
        pdf_check = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(pdf_check)
        pdf_check.close()

        st.info(f"Total pages detected: {total_pages}")

    except Exception:
        st.error("This file could not be opened as a valid PDF.")
        st.stop()

    if st.button("🚀 Extract Text"):
        try:
            with st.spinner("Extracting text. Please wait..."):

                if extraction_mode == "Searchable PDF Only":
                    extracted_text = extract_searchable_pdf_text(pdf_bytes)
                    extraction_message = "Direct text extraction completed."

                elif extraction_mode == "Scanned PDF OCR Only":
                    if not tesseract_available:
                        st.error(
                            "Tesseract OCR is not installed on the server. "
                            "Add packages.txt in GitHub and redeploy the app."
                        )
                        st.stop()

                    if not selected_language_available:
                        st.error(
                            f"The selected OCR language '{ocr_language}' is not installed. "
                            "Add it in packages.txt or select another language."
                        )
                        st.stop()

                    extracted_text = extract_scanned_pdf_text(
                        pdf_bytes,
                        dpi=ocr_dpi,
                        language=ocr_language
                    )
                    extraction_message = "OCR extraction completed."

                else:
                    # Auto Detect
                    # First check whether the PDF already has searchable text.
                    test_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    test_text = ""

                    for page in test_doc:
                        test_text += page.get_text("text").strip()

                    test_doc.close()

                    if len(test_text) > 30:
                        extracted_text = extract_searchable_pdf_text(pdf_bytes)
                        extraction_message = "Searchable PDF detected. Direct extraction completed."
                    else:
                        if not tesseract_available:
                            st.error(
                                "This PDF seems scanned, but Tesseract OCR is not installed on the server. "
                                "Add packages.txt in GitHub and redeploy the app."
                            )
                            st.stop()

                        if not selected_language_available:
                            st.error(
                                f"This PDF needs OCR, but OCR language '{ocr_language}' is not installed."
                            )
                            st.stop()

                        extracted_text = extract_scanned_pdf_text(
                            pdf_bytes,
                            dpi=ocr_dpi,
                            language=ocr_language
                        )
                        extraction_message = "Scanned PDF detected. OCR extraction completed."

                extracted_text = clean_text(extracted_text)

            st.success(extraction_message)

            if not extracted_text:
                st.warning(
                    "No text was extracted. The PDF may be blurred, handwritten, damaged, "
                    "or the OCR language may not be installed."
                )

            else:
                st.subheader("📌 Extracted Text Preview")

                st.text_area(
                    "Preview of extracted text",
                    extracted_text[:10000],
                    height=350
                )

                st.write(f"Total extracted characters: {len(extracted_text)}")

                if output_type == "TXT File":
                    txt_file = create_txt_file(extracted_text)

                    st.download_button(
                        label="⬇️ Download TXT File",
                        data=txt_file,
                        file_name="extracted_pdf_text.txt",
                        mime="text/plain"
                    )

                else:
                    word_file = create_word_file(extracted_text)

                    st.download_button(
                        label="⬇️ Download Word File",
                        data=word_file,
                        file_name="extracted_pdf_text.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        except Exception as e:
            st.error("An error occurred while extracting text.")
            st.exception(e)

else:
    st.info("Please upload a PDF file to start extraction.")