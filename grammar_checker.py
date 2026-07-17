from __future__ import annotations

import html
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

import fitz
import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX


# =========================================================
# STREAMLIT PAGE SETTINGS
# =========================================================

st.set_page_config(
    page_title="Spelling and Grammar Checker",
    page_icon="📝",
    layout="wide",
)


# =========================================================
# APPLICATION SETTINGS
# =========================================================

LANGUAGES = {
    "English (United States)": "en-US",
    "English (United Kingdom)": "en-GB",
    "English (Canada)": "en-CA",
    "English (Australia)": "en-AU",
    "Automatic detection": "auto",
}

DEFAULT_API_URL = "https://api.languagetool.org/v2/check"

MAX_CHUNK_SIZE = 9000
REQUEST_TIMEOUT = 60


# =========================================================
# ISSUE DATA STRUCTURE
# =========================================================

@dataclass(frozen=True)
class Issue:
    offset: int
    length: int
    message: str
    replacements: tuple[str, ...]
    category: str
    rule_id: str
    sentence: str

    @property
    def end(self) -> int:
        return self.offset + self.length

    @property
    def first_replacement(self) -> str:
        if self.replacements:
            return self.replacements[0]

        return ""


# =========================================================
# DOCUMENT TEXT EXTRACTION
# =========================================================

def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract text from paragraphs and tables in a DOCX file.
    """

    document = Document(io.BytesIO(file_bytes))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        blocks.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_values: list[str] = []

            for cell in row.cells:
                clean_text = re.sub(
                    r"\s*\n\s*",
                    " ",
                    cell.text,
                ).strip()

                row_values.append(clean_text)

            blocks.append("\t".join(row_values))

    return "\n".join(blocks).strip()


def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract selectable text from a PDF file.
    """

    pages: list[str] = []

    with fitz.open(
        stream=file_bytes,
        filetype="pdf",
    ) as pdf:

        for page_number, page in enumerate(
            pdf,
            start=1,
        ):
            page_text = page.get_text("text").strip()

            if page_text:
                pages.append(
                    f"[Page {page_number}]\n{page_text}"
                )

    return "\n\n".join(pages).strip()


def extract_uploaded_text(
    uploaded_file: Any,
) -> str:
    """
    Choose the correct extraction method.
    """

    extension = Path(
        uploaded_file.name
    ).suffix.lower()

    file_bytes = uploaded_file.getvalue()

    if extension == ".docx":
        return extract_docx_text(file_bytes)

    if extension == ".pdf":
        return extract_pdf_text(file_bytes)

    raise ValueError(
        "Only DOCX and PDF files are supported."
    )


# =========================================================
# DOCUMENT CHUNKING
# =========================================================

def split_text_with_offsets(
    text: str,
    max_chars: int = MAX_CHUNK_SIZE,
) -> Iterable[tuple[int, str]]:
    """
    Split large documents into smaller parts while preserving
    the original character positions.
    """

    start = 0
    total_length = len(text)

    while start < total_length:
        proposed_end = min(
            start + max_chars,
            total_length,
        )

        end = proposed_end

        if proposed_end < total_length:
            search_start = start + int(
                max_chars * 0.60
            )

            break_positions = [
                text.rfind(
                    "\n\n",
                    search_start,
                    proposed_end,
                ),
                text.rfind(
                    "\n",
                    search_start,
                    proposed_end,
                ),
                text.rfind(
                    ". ",
                    search_start,
                    proposed_end,
                ),
                text.rfind(
                    "? ",
                    search_start,
                    proposed_end,
                ),
                text.rfind(
                    "! ",
                    search_start,
                    proposed_end,
                ),
                text.rfind(
                    " ",
                    search_start,
                    proposed_end,
                ),
            ]

            best_break = max(break_positions)

            if best_break > start:
                break_characters = text[
                    best_break:best_break + 2
                ]

                if break_characters in {
                    "\n\n",
                    ". ",
                    "? ",
                    "! ",
                }:
                    end = best_break + 2
                else:
                    end = best_break + 1

        if end <= start:
            end = proposed_end

        yield start, text[start:end]

        start = end


# =========================================================
# LANGUAGETOOL FUNCTIONS
# =========================================================

def check_text_chunk(
    text: str,
    language: str,
    api_url: str,
) -> list[dict[str, Any]]:
    """
    Send one section of text to LanguageTool.
    """

    response = requests.post(
        api_url,
        data={
            "text": text,
            "language": language,
            "enabledOnly": "false",
        },
        timeout=REQUEST_TIMEOUT,
    )

    response.raise_for_status()

    response_data = response.json()

    return response_data.get("matches", [])


def convert_match_to_issue(
    match: dict[str, Any],
    chunk_start: int,
) -> Issue:
    """
    Convert a LanguageTool match into an Issue object.
    """

    rule = match.get("rule", {})
    category = rule.get("category", {})

    replacements = tuple(
        replacement.get("value", "")
        for replacement in match.get(
            "replacements",
            [],
        )
        if replacement.get("value")
    )

    sentence = str(
        match.get("sentence")
        or match.get(
            "context",
            {},
        ).get("text", "")
    ).strip()

    return Issue(
        offset=(
            chunk_start
            + int(match.get("offset", 0))
        ),
        length=int(match.get("length", 0)),
        message=str(
            match.get(
                "message",
                "Possible spelling or grammar error",
            )
        ),
        replacements=replacements,
        category=str(
            category.get(
                "name",
                "Other",
            )
        ),
        rule_id=str(
            rule.get("id", "")
        ),
        sentence=sentence,
    )


def check_complete_document(
    text: str,
    language: str,
    api_url: str,
) -> list[Issue]:
    """
    Check the complete document section by section.
    """

    chunks = list(
        split_text_with_offsets(text)
    )

    detected_issues: list[Issue] = []

    progress_bar = st.progress(
        0,
        text="Preparing document...",
    )

    for chunk_number, (
        chunk_start,
        chunk_text,
    ) in enumerate(
        chunks,
        start=1,
    ):

        progress_bar.progress(
            chunk_number / len(chunks),
            text=(
                f"Checking section {chunk_number} "
                f"of {len(chunks)}..."
            ),
        )

        matches = check_text_chunk(
            text=chunk_text,
            language=language,
            api_url=api_url,
        )

        for match in matches:
            issue = convert_match_to_issue(
                match=match,
                chunk_start=chunk_start,
            )

            if (
                issue.length > 0
                and issue.offset >= 0
                and issue.end <= len(text)
            ):
                detected_issues.append(issue)

    progress_bar.empty()

    unique_issues: dict[
        tuple[int, int, str],
        Issue,
    ] = {}

    for issue in detected_issues:
        unique_key = (
            issue.offset,
            issue.length,
            issue.rule_id,
        )

        unique_issues[unique_key] = issue

    return sorted(
        unique_issues.values(),
        key=lambda issue: (
            issue.offset,
            issue.length,
        ),
    )


# =========================================================
# CORRECTION FUNCTIONS
# =========================================================

def select_non_overlapping_issues(
    issues: list[Issue],
) -> list[Issue]:
    """
    Remove overlapping issues so that replacements are applied
    safely without damaging the text.
    """

    selected_issues: list[Issue] = []
    last_end_position = -1

    sorted_issues = sorted(
        issues,
        key=lambda issue: (
            issue.offset,
            -issue.length,
        ),
    )

    for issue in sorted_issues:
        if issue.offset >= last_end_position:
            selected_issues.append(issue)
            last_end_position = issue.end

    return selected_issues


def apply_all_corrections(
    text: str,
    issues: list[Issue],
) -> tuple[str, int]:
    """
    Apply the first available replacement for every correctable
    issue.

    Corrections are applied from the end of the document towards
    the beginning so character positions remain valid.
    """

    correctable_issues = [
        issue
        for issue in select_non_overlapping_issues(
            issues
        )
        if issue.replacements
    ]

    corrected_text = text

    for issue in sorted(
        correctable_issues,
        key=lambda issue: issue.offset,
        reverse=True,
    ):
        corrected_text = (
            corrected_text[:issue.offset]
            + issue.first_replacement
            + corrected_text[issue.end:]
        )

    return corrected_text, len(correctable_issues)


# =========================================================
# HTML ERROR PREVIEW
# =========================================================

def create_highlighted_html(
    text: str,
    issues: list[Issue],
) -> str:
    """
    Create an HTML preview with highlighted errors.
    """

    output_parts: list[str] = []
    cursor = 0

    for issue in select_non_overlapping_issues(
        issues
    ):
        normal_text = text[
            cursor:issue.offset
        ]

        output_parts.append(
            html.escape(normal_text)
        )

        incorrect_text = html.escape(
            text[issue.offset:issue.end]
        )

        if issue.replacements:
            suggested_text = ", ".join(
                issue.replacements[:5]
            )
        else:
            suggested_text = (
                "No automatic correction available"
            )

        tooltip = html.escape(
            (
                f"{issue.message} "
                f"Suggested: {suggested_text}"
            ),
            quote=True,
        )

        output_parts.append(
            f'<span class="error-text" '
            f'title="{tooltip}">'
            f"{incorrect_text}"
            f"</span>"
        )

        cursor = issue.end

    output_parts.append(
        html.escape(text[cursor:])
    )

    highlighted_content = "".join(
        output_parts
    ).replace(
        "\n",
        "<br>",
    )

    return f"""
    <style>
        .document-preview {{
            max-height: 650px;
            overflow-y: auto;
            padding: 22px;
            border: 1px solid rgba(128, 128, 128, 0.35);
            border-radius: 10px;
            line-height: 1.8;
            font-family: Arial, sans-serif;
            background-color: rgba(128, 128, 128, 0.03);
        }}

        .error-text {{
            background-color: rgba(255, 215, 0, 0.55);
            text-decoration-line: underline;
            text-decoration-style: wavy;
            text-decoration-color: red;
            text-decoration-thickness: 2px;
            cursor: help;
        }}
    </style>

    <div class="document-preview">
        {highlighted_content}
    </div>
    """


# =========================================================
# CORRECTION COMPARISON
# =========================================================

def create_correction_comparison_html(
    text: str,
    issues: list[Issue],
) -> str:
    """
    Show original errors and automatic corrections.
    """

    output_parts: list[str] = []
    cursor = 0

    correctable_issues = [
        issue
        for issue in select_non_overlapping_issues(
            issues
        )
        if issue.replacements
    ]

    for issue in correctable_issues:
        output_parts.append(
            html.escape(
                text[cursor:issue.offset]
            )
        )

        original_text = html.escape(
            text[issue.offset:issue.end]
        )

        corrected_text = html.escape(
            issue.first_replacement
        )

        output_parts.append(
            f'<span class="original-error">'
            f"{original_text}"
            f"</span>"
            f'<span class="change-arrow"> → </span>'
            f'<span class="corrected-value">'
            f"{corrected_text}"
            f"</span>"
        )

        cursor = issue.end

    output_parts.append(
        html.escape(text[cursor:])
    )

    comparison_content = "".join(
        output_parts
    ).replace(
        "\n",
        "<br>",
    )

    return f"""
    <style>
        .comparison-preview {{
            max-height: 650px;
            overflow-y: auto;
            padding: 22px;
            border: 1px solid rgba(128, 128, 128, 0.35);
            border-radius: 10px;
            line-height: 1.9;
            font-family: Arial, sans-serif;
        }}

        .original-error {{
            background-color: rgba(255, 80, 80, 0.25);
            color: #a00000;
            text-decoration: line-through;
            padding: 2px 5px;
            border-radius: 4px;
        }}

        .corrected-value {{
            background-color: rgba(40, 180, 80, 0.25);
            color: #126b2f;
            font-weight: bold;
            padding: 2px 5px;
            border-radius: 4px;
        }}

        .change-arrow {{
            font-weight: bold;
            padding: 0 4px;
        }}
    </style>

    <div class="comparison-preview">
        {comparison_content}
    </div>
    """


# =========================================================
# WORD DOCUMENT CREATION
# =========================================================

def add_text_with_breaks(
    paragraph: Any,
    text: str,
) -> None:
    """
    Add multiline text to a Word paragraph.
    """

    lines = text.split("\n")

    for line_number, line in enumerate(lines):
        if line_number > 0:
            paragraph.add_run().add_break()

        paragraph.add_run(line)


def create_corrected_docx(
    text: str,
) -> bytes:
    """
    Create a downloadable corrected Word document.
    """

    document = Document()

    document.add_heading(
        "Corrected Document",
        level=1,
    )

    blocks = re.split(
        r"\n{2,}",
        text,
    )

    for block in blocks:
        paragraph = document.add_paragraph()

        add_text_with_breaks(
            paragraph,
            block,
        )

    output_file = io.BytesIO()
    document.save(output_file)

    return output_file.getvalue()


def create_highlighted_docx(
    text: str,
    issues: list[Issue],
) -> bytes:
    """
    Create a Word document with detected issues highlighted.
    """

    document = Document()

    document.add_heading(
        "Spelling and Grammar Review",
        level=1,
    )

    document.add_paragraph(
        "Highlighted text indicates a possible "
        "spelling, grammar, punctuation, or style issue."
    )

    paragraph = document.add_paragraph()

    cursor = 0

    for issue in select_non_overlapping_issues(
        issues
    ):
        normal_text = text[
            cursor:issue.offset
        ]

        error_text = text[
            issue.offset:issue.end
        ]

        paragraph.add_run(normal_text)

        error_run = paragraph.add_run(
            error_text
        )

        error_run.font.highlight_color = (
            WD_COLOR_INDEX.YELLOW
        )

        cursor = issue.end

    paragraph.add_run(
        text[cursor:]
    )

    document.add_page_break()

    document.add_heading(
        "Issue Report",
        level=1,
    )

    for issue_number, issue in enumerate(
        issues,
        start=1,
    ):
        original_text = text[
            issue.offset:issue.end
        ]

        if issue.replacements:
            suggestions = ", ".join(
                issue.replacements[:5]
            )
        else:
            suggestions = (
                "Manual review required"
            )

        document.add_paragraph(
            f"{issue_number}. {issue.category}",
            style="List Number",
        )

        document.add_paragraph(
            f"Original text: {original_text}"
        )

        document.add_paragraph(
            f"Explanation: {issue.message}"
        )

        document.add_paragraph(
            f"Suggested correction: {suggestions}"
        )

        if issue.sentence:
            document.add_paragraph(
                f"Context: {issue.sentence}"
            )

    output_file = io.BytesIO()
    document.save(output_file)

    return output_file.getvalue()


# =========================================================
# ISSUE REPORT
# =========================================================

def create_issue_dataframe(
    text: str,
    issues: list[Issue],
) -> pd.DataFrame:
    """
    Convert detected issues to a table.
    """

    table_rows: list[dict[str, Any]] = []

    for issue_number, issue in enumerate(
        issues,
        start=1,
    ):
        if issue.replacements:
            main_correction = (
                issue.first_replacement
            )

            other_corrections = ", ".join(
                issue.replacements[1:5]
            )
        else:
            main_correction = (
                "Manual review required"
            )

            other_corrections = ""

        table_rows.append(
            {
                "No.": issue_number,
                "Error": text[
                    issue.offset:issue.end
                ],
                "Correction": main_correction,
                "Other suggestions": (
                    other_corrections
                ),
                "Category": issue.category,
                "Explanation": issue.message,
                "Context": issue.sentence,
            }
        )

    return pd.DataFrame(table_rows)


# =========================================================
# SESSION STATE MANAGEMENT
# =========================================================

def reset_session_for_new_file(
    file_signature: str | None,
) -> None:
    """
    Remove previous results when a different file is uploaded.
    """

    previous_signature = st.session_state.get(
        "file_signature"
    )

    if previous_signature != file_signature:
        st.session_state[
            "file_signature"
        ] = file_signature

        keys_to_remove = [
            "source_text",
            "issues",
            "original_filename",
            "corrected_editor",
            "corrections_applied",
            "applied_correction_count",
        ]

        for key in keys_to_remove:
            st.session_state.pop(
                key,
                None,
            )


# =========================================================
# MAIN STREAMLIT APPLICATION
# =========================================================

def main() -> None:
    st.title(
        "📝 Spelling and Grammar Checker"
    )

    st.write(
        "Upload a Word document or text-based PDF, "
        "check it for spelling and grammar errors, "
        "and then use the correction button to apply "
        "the available corrections."
    )

    # -----------------------------------------------------
    # SIDEBAR
    # -----------------------------------------------------

    with st.sidebar:
        st.header("Settings")

        language_name = st.selectbox(
            "Document language",
            options=list(
                LANGUAGES.keys()
            ),
            index=1,
        )

        language_code = LANGUAGES[
            language_name
        ]

        api_url = st.text_input(
            "LanguageTool API URL",
            value=DEFAULT_API_URL,
            help=(
                "Use the public API address or the "
                "address of your private LanguageTool server."
            ),
        ).strip()

        st.info(
            "First check the document. After checking, "
            "use the Apply all corrections button."
        )

    # -----------------------------------------------------
    # FILE UPLOAD
    # -----------------------------------------------------

    uploaded_file = st.file_uploader(
        "Upload a DOCX or PDF document",
        type=["docx", "pdf"],
        accept_multiple_files=False,
    )

    file_signature = None

    if uploaded_file is not None:
        file_signature = (
            f"{uploaded_file.name}:"
            f"{uploaded_file.size}"
        )

    reset_session_for_new_file(
        file_signature
    )

    if uploaded_file is None:
        st.info(
            "Upload a .docx or .pdf file to begin."
        )
        return

    st.success(
        f"Selected file: {uploaded_file.name}"
    )

    # -----------------------------------------------------
    # CHECK BUTTON
    # -----------------------------------------------------

    check_document_button = st.button(
        "🔍 Check document",
        type="primary",
        use_container_width=True,
    )

    if check_document_button:
        try:
            with st.spinner(
                "Extracting text..."
            ):
                source_text = extract_uploaded_text(
                    uploaded_file
                )

            if not source_text.strip():
                st.error(
                    "No selectable text was found. "
                    "The PDF may be scanned and must "
                    "be processed with OCR first."
                )
                return

            with st.spinner(
                "Checking spelling and grammar..."
            ):
                issues = check_complete_document(
                    text=source_text,
                    language=language_code,
                    api_url=api_url,
                )

            st.session_state[
                "source_text"
            ] = source_text

            st.session_state[
                "issues"
            ] = issues

            st.session_state[
                "original_filename"
            ] = uploaded_file.name

            st.session_state[
                "corrected_editor"
            ] = source_text

            st.session_state[
                "corrections_applied"
            ] = False

            st.session_state[
                "applied_correction_count"
            ] = 0

            st.success(
                "Document checking completed."
            )

        except requests.Timeout:
            st.error(
                "The grammar server took too long "
                "to respond."
            )
            return

        except requests.RequestException as error:
            st.error(
                f"Grammar checking failed: {error}"
            )
            return

        except Exception as error:
            st.error(
                f"The document could not be processed: {error}"
            )
            return

    # -----------------------------------------------------
    # WAIT UNTIL DOCUMENT HAS BEEN CHECKED
    # -----------------------------------------------------

    if "source_text" not in st.session_state:
        return

    source_text: str = st.session_state[
        "source_text"
    ]

    issues: list[Issue] = st.session_state[
        "issues"
    ]

    original_filename: str = st.session_state[
        "original_filename"
    ]

    base_filename = Path(
        original_filename
    ).stem

    correctable_issue_count = sum(
        1
        for issue in select_non_overlapping_issues(
            issues
        )
        if issue.replacements
    )

    manual_review_count = sum(
        1
        for issue in issues
        if not issue.replacements
    )

    # -----------------------------------------------------
    # RESULT METRICS
    # -----------------------------------------------------

    metric_1, metric_2, metric_3, metric_4 = (
        st.columns(4)
    )

    metric_1.metric(
        "Words",
        f"{len(source_text.split()):,}",
    )

    metric_2.metric(
        "Detected issues",
        f"{len(issues):,}",
    )

    metric_3.metric(
        "Available corrections",
        f"{correctable_issue_count:,}",
    )

    metric_4.metric(
        "Manual review",
        f"{manual_review_count:,}",
    )

    if not issues:
        st.success(
            "No spelling or grammar errors were detected."
        )
        return

    # -----------------------------------------------------
    # CORRECTION CONTROL BUTTONS
    # -----------------------------------------------------

    st.subheader(
        "Correction controls"
    )

    correction_column, reset_column = st.columns(2)

    with correction_column:
        apply_corrections_button = st.button(
            "✅ Apply all corrections",
            type="primary",
            use_container_width=True,
            disabled=(
                correctable_issue_count == 0
            ),
        )

    with reset_column:
        reset_corrections_button = st.button(
            "↩️ Reset to original",
            use_container_width=True,
        )

    if apply_corrections_button:
        corrected_text, correction_count = (
            apply_all_corrections(
                text=source_text,
                issues=issues,
            )
        )

        st.session_state[
            "corrected_editor"
        ] = corrected_text

        st.session_state[
            "corrections_applied"
        ] = True

        st.session_state[
            "applied_correction_count"
        ] = correction_count

        st.rerun()

    if reset_corrections_button:
        st.session_state[
            "corrected_editor"
        ] = source_text

        st.session_state[
            "corrections_applied"
        ] = False

        st.session_state[
            "applied_correction_count"
        ] = 0

        st.rerun()

    if st.session_state.get(
        "corrections_applied",
        False,
    ):
        st.success(
            f"{st.session_state.get('applied_correction_count', 0)} "
            "corrections have been applied."
        )
    else:
        st.warning(
            "Corrections have not been applied yet. "
            "Click Apply all corrections."
        )

    # -----------------------------------------------------
    # RESULT TABS
    # -----------------------------------------------------

    (
        highlighted_tab,
        comparison_tab,
        corrected_tab,
        report_tab,
    ) = st.tabs(
        [
            "Highlighted errors",
            "Correction comparison",
            "Corrected document",
            "Issue report",
        ]
    )

    # -----------------------------------------------------
    # HIGHLIGHTED ERRORS TAB
    # -----------------------------------------------------

    with highlighted_tab:
        st.markdown(
            create_highlighted_html(
                text=source_text,
                issues=issues,
            ),
            unsafe_allow_html=True,
        )

        highlighted_docx = create_highlighted_docx(
            text=source_text,
            issues=issues,
        )

        st.download_button(
            "Download highlighted DOCX",
            data=highlighted_docx,
            file_name=(
                f"{base_filename}_highlighted.docx"
            ),
            mime=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),
            use_container_width=True,
        )

    # -----------------------------------------------------
    # COMPARISON TAB
    # -----------------------------------------------------

    with comparison_tab:
        if correctable_issue_count > 0:
            st.caption(
                "Red crossed-out text is the original. "
                "Green text is the suggested correction."
            )

            st.markdown(
                create_correction_comparison_html(
                    text=source_text,
                    issues=issues,
                ),
                unsafe_allow_html=True,
            )
        else:
            st.info(
                "No automatic correction suggestions "
                "were available."
            )

    # -----------------------------------------------------
    # CORRECTED DOCUMENT TAB
    # -----------------------------------------------------

    with corrected_tab:
        if not st.session_state.get(
            "corrections_applied",
            False,
        ):
            st.warning(
                "The text below is still the original document. "
                "Click Apply all corrections above."
            )
        else:
            st.success(
                "Automatic corrections have been applied. "
                "Review the text before downloading it."
            )

        edited_text = st.text_area(
            "Review and edit the corrected text",
            key="corrected_editor",
            height=600,
        )

        corrected_docx = create_corrected_docx(
            edited_text
        )

        download_column_1, download_column_2 = (
            st.columns(2)
        )

        with download_column_1:
            st.download_button(
                "Download corrected DOCX",
                data=corrected_docx,
                file_name=(
                    f"{base_filename}_corrected.docx"
                ),
                mime=(
                    "application/vnd.openxmlformats-"
                    "officedocument.wordprocessingml.document"
                ),
                use_container_width=True,
            )

        with download_column_2:
            st.download_button(
                "Download corrected TXT",
                data=edited_text.encode("utf-8"),
                file_name=(
                    f"{base_filename}_corrected.txt"
                ),
                mime="text/plain",
                use_container_width=True,
            )

    # -----------------------------------------------------
    # ISSUE REPORT TAB
    # -----------------------------------------------------

    with report_tab:
        issue_dataframe = create_issue_dataframe(
            text=source_text,
            issues=issues,
        )

        st.dataframe(
            issue_dataframe,
            use_container_width=True,
            hide_index=True,
        )

        csv_data = issue_dataframe.to_csv(
            index=False
        ).encode("utf-8-sig")

        st.download_button(
            "Download issue report as CSV",
            data=csv_data,
            file_name=(
                f"{base_filename}_issue_report.csv"
            ),
            mime="text/csv",
            use_container_width=True,
        )


if __name__ == "__main__":
    main()